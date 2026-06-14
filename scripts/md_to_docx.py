#!/usr/bin/env python3
"""Convert SYSTEM_ARCHITECTURE_ANALYSIS.md to Word with basic formatting."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
MD_FILE = ROOT / "SYSTEM_ARCHITECTURE_ANALYSIS.md"
OUTPUT = ROOT / "docs" / "Enterprise_AI_Platform_Architecture_Analysis.docx"


def parse_table_lines(lines):
    rows = []
    for line in lines:
        if not line.strip().startswith("|"):
            break
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(set(c) <= {"-", ":"} for c in cells):
            continue  # separator row
        rows.append(cells)
    return rows


def add_md_table(doc, rows):
    if not rows:
        return
    headers = rows[0]
    body = rows[1:] if len(rows) > 1 else []
    table = doc.add_table(rows=1 + len(body), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for ri, row in enumerate(body):
        for ci in range(len(headers)):
            text = row[ci] if ci < len(row) else ""
            table.rows[ri + 1].cells[ci].text = text
    doc.add_paragraph()


def add_formatted_paragraph(doc, text):
    """Handle **bold** inline."""
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part:
            p.add_run(part)


def convert(md_path: Path, out_path: Path):
    content = md_path.read_text(encoding="utf-8")
    lines = content.splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Cover page
    title = doc.add_heading("Enterprise AI Platform", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_heading("Complete System Architecture Analysis", level=1)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Reverse-engineering document for interviews, production ownership, scaling, and architectural defense. "
        "Covers Phases 1–14: architecture, APIs, databases, events, cache, security, observability, scalability, "
        "failures, interviews, and critique."
    )
    doc.add_paragraph(f"Source: {md_path.name}")
    doc.add_page_break()

    i = 0
    in_code = False
    code_buf = []
    list_buf = []

    def flush_list():
        nonlocal list_buf
        for item in list_buf:
            text = re.sub(r"^\d+\.\s+", "", item)
            text = re.sub(r"^[-*]\s+", "", text)
            doc.add_paragraph(text, style="List Bullet")
        list_buf = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_buf = []
                in_code = False
            else:
                flush_list()
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if line.strip().startswith("|"):
            flush_list()
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_md_table(doc, parse_table_lines(table_lines))
            continue

        if re.match(r"^#{1,6}\s+", line):
            flush_list()
            level = len(line) - len(line.lstrip("#"))
            text = line.lstrip("#").strip()
            # Skip duplicate TOC entries at depth 1 if too nested
            doc.add_heading(text, level=min(level, 3))
            i += 1
            continue

        if re.match(r"^[-*]\s+", line) or re.match(r"^\d+\.\s+", line):
            list_buf.append(line.strip())
            i += 1
            continue

        if line.strip() == "---":
            flush_list()
            doc.add_page_break()
            i += 1
            continue

        if not line.strip():
            flush_list()
            i += 1
            continue

        flush_list()
        # Skip mermaid blocks marker lines
        if line.strip().startswith("```mermaid"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            i += 1
            doc.add_paragraph("[Diagram: see markdown source for Mermaid sequence/flowchart]")
            continue

        add_formatted_paragraph(doc, line.strip())
        i += 1

    flush_list()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"Wrote {out_path} ({out_path.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    convert(MD_FILE, OUTPUT)
