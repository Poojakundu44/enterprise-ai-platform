#!/usr/bin/env python3
"""Generate Enterprise AI Platform architecture analysis Word document."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "docs" / "Enterprise_AI_Platform_Architecture_Analysis.docx"


def add_title(doc, text, level=0):
    if level == 0:
        p = doc.add_heading(text, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_heading(text, level=min(level, 3))


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def build_document():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_title(doc, "Enterprise AI Platform")
    add_title(doc, "Complete Architecture Analysis & Reverse Engineering", level=0)
    add_para(
        doc,
        "Document version: Phase 1 MVP (0.0.1-SNAPSHOT) | Generated for senior/staff/principal interviews, "
        "production ownership, and scaling to millions of users.",
    )
    doc.add_page_break()

    # --- PHASE 1 ---
    add_title(doc, "Phase 1: High-Level Architecture", level=1)

    add_title(doc, "1. System Overview", level=2)
    add_para(
        doc,
        "The Enterprise AI Platform is a cloud-native, microservices-oriented system that lets enterprise employees "
        "securely search, chat with, and automate work across organizational knowledge (Confluence, Jira, SharePoint, "
        "PDFs, incident reports) using RAG, AI agents, and workflow automation. "
        "Phase 1 MVP delivers: JWT auth, manual document upload, ingestion/chunking, dev embeddings, semantic search, "
        "and RAG chat (stub or OpenAI). Agent, audit, Kafka, OIDC, and production hardening are planned.",
    )

    add_title(doc, "2. Business Objective", level=2)
    add_table(
        doc,
        ["Driver", "How the platform addresses it"],
        [
            ["Knowledge fragmentation", "Unified search and Q&A across many systems"],
            ["Unsafe ad-hoc LLM use", "Central gateway, auth, audit, policy enforcement"],
            ["Manual repetitive work", "Agent workflows with enterprise tool integrations"],
            ["Compliance", "Immutable audit trail, access control on sources"],
            ["Time-to-answer", "RAG with grounded citations from internal docs"],
        ],
    )

    add_title(doc, "3. Core User Journey (MVP)", level=2)
    add_bullets(
        doc,
        [
            "Employee registers or logs in via gateway → receives JWT (HS256, 1h default).",
            "Employee uploads a text document via POST /api/v1/connectors/documents.",
            "Connector forwards to ingestion → chunks (800 chars, 100 overlap) → embedding index.",
            "Employee asks a question via POST /api/v1/rag/chat → search retrieves top-K chunks → LLM answers with citations.",
        ],
    )

    add_title(doc, "4. End-to-End Request Flow", level=2)
    add_code(
        doc,
        "Client → Load Balancer (prod) → gateway-service:8080 (JWT validate, X-User-Email)\n"
        "  → auth | connector | search | rag (domain routes)\n"
        "connector → ingestion:8083 → embedding:8084 (sync WebClient)\n"
        "rag → search:8085 → embedding:8084\n"
        "rag → OpenAI API (optional, if OPENAI_API_KEY set)",
    )

    add_title(doc, "5. Architecture Diagram (ASCII)", level=2)
    add_code(
        doc,
        """
┌─────────────┐     ┌──────────────────┐     ┌─────────────────────────────────────────┐
│  Client /   │     │ gateway-service  │     │           Microservices (8081-8089)      │
│  Frontend   │────▶│  :8080           │────▶│ auth │ connector │ ingestion │ embedding │
│  (planned)  │     │ JWT + routes     │     │ search │ rag │ agent* │ audit* │ notif*   │
└─────────────┘     └──────────────────┘     └───────────┬─────────────────────────────┘
                                                         │
                    ┌────────────────────────────────────┼────────────────────┐
                    ▼                                    ▼                    ▼
            ┌───────────────┐                  ┌─────────────────┐    ┌──────────────┐
            │ PostgreSQL 16 │                  │ OpenAI (opt.)   │    │ Prometheus / │
            │ auth, ingest, │                  │ chat completions│    │ Grafana*     │
            │ embed, audit* │                  └─────────────────┘    └──────────────┘
            └───────────────┘
* = scaffold or partial
""".strip(),
    )

    add_title(doc, "6. Deployment Architecture", level=2)
    add_para(
        doc,
        "Development: single Postgres container (infra/docker-compose.yml), 7+ JVM processes on localhost, "
        "hardcoded gateway routes. Production target: Kubernetes with Ingress → gateway Deployment (HPA), "
        "one Deployment per service, managed Postgres (or RDS), secrets via Vault, service mesh/mTLS for east-west, "
        "vector store (pgvector/Pinecone), object storage for raw docs, Kafka for async ingestion.",
    )

    add_title(doc, "7. Data Flow Diagram", level=2)
    add_code(
        doc,
        "Upload: DocumentUploadRequest → connector → ingestion (documents + document_chunks tables)\n"
        "         → embedding (chunk_embeddings.embedding_json)\n"
        "Query:  ChatRequest → rag → search → embedding (cosine on all rows — MVP only)\n"
        "         → LlmAnswerService → ChatResponse (answer + citations + model)",
    )

    add_title(doc, "8. Service Dependency Diagram", level=2)
    add_table(
        doc,
        ["Service", "Depends on", "Database"],
        [
            ["gateway", "auth (route), all backends", "—"],
            ["auth", "Postgres", "enterprise_ai_auth"],
            ["connector", "ingestion (HTTP)", "—"],
            ["ingestion", "embedding (HTTP)", "enterprise_ai_ingestion"],
            ["embedding", "Postgres", "enterprise_ai_embedding"],
            ["search", "embedding (HTTP)", "—"],
            ["rag", "search (HTTP), OpenAI (opt.)", "—"],
            ["agent, notification", "— (scaffold)", "—"],
            ["audit", "Postgres (empty schema)", "enterprise_ai_audit"],
        ],
    )

    add_title(doc, "Architecture Decision Rationale", level=2)
    add_bullets(
        doc,
        [
            "Microservices chosen: independent scaling (ingestion vs RAG), blast-radius isolation, team ownership. "
            "Rejected monolith: noisy deploys, cannot scale embedding workers separately.",
            "Monorepo Maven: shared DTOs (platform-common), atomic version bumps across 10 services.",
            "Spring Cloud Gateway: reactive edge, central JWT, future rate limits and SSE for streaming chat.",
            "Sync REST for MVP: simplest vertical slice; Kafka planned before production ingestion load.",
            "Per-service Postgres DB: auth vs ingestion vs embedding isolation; audit DB provisioned early.",
            "Security: JWT validated only at gateway (MVP trade-off); production needs mTLS + service tokens.",
            "Cost: microservices ops overhead early; LLM token cost dominates RAG at scale, not JVM CPU.",
        ],
    )
    doc.add_page_break()

    # --- PHASE 2 ---
    add_title(doc, "Phase 2: Folder Structure Analysis", level=1)
    folders = [
        ("platform-common/", "Shared DTOs, ApiResponse envelope, PlatformConstants. Without it: duplicated API contracts."),
        ("platform-service-parent/", "Maven BOM for servlet services. Gateway intentionally excludes (WebFlux)."),
        ("gateway-service/", "North-south entry, JWT, routing. Removed: no single public API."),
        ("auth-service/", "Identity, JWT issuance, users table. Merged into RAG would violate security boundary."),
        ("connector-service/", "External system integrations; isolates third-party API volatility."),
        ("ingestion-service/", "CPU/IO heavy chunking; must not block search latency path."),
        ("embedding-service/", "Vector generation and index; scales on GPU/batch, not chat QPS."),
        ("search-service/", "Sub-100ms retrieval path separate from LLM orchestration."),
        ("rag-service/", "Highest-value API; isolates prompt logic and LLM spend controls."),
        ("agent-service/", "Long-running stateful workflows — scaffold only."),
        ("audit-service/", "Append-only compliance log — separate DB from auth."),
        ("notification-service/", "Async side effects decoupled from RAG critical path."),
        ("frontend/", "UI placeholder; separates deploy cadence from backend."),
        ("observability/", "Prometheus + Grafana local; scrape config ahead of micrometer-prometheus dep."),
        ("infra/", "Postgres compose + init-databases.sql."),
        ("docs/", "ARCHITECTURE.md, MVP_RUNBOOK.md, this document."),
    ]
    for folder, desc in folders:
        add_para(doc, f"{folder}", bold=True)
        add_para(doc, desc)
    doc.add_page_break()

    # --- PHASE 3 ---
    add_title(doc, "Phase 3: Service-by-Service Deep Dive", level=1)
    services_detail = [
        (
            "gateway-service (8080)",
            "Business: Single front door; enforce authn/z, rate limits, routing.",
            "Technical: Spring Cloud Gateway routes; GatewaySecurityConfig (JWT HS256); UserEmailGatewayFilter.",
            "Why separate: Edge policy vs domain logic; reactive stack for future SSE.",
            "Scaling: Stateless horizontal scale behind LB. Failure: backend down → 503.",
        ),
        (
            "auth-service (8081)",
            "Business: Login, register, future OIDC/SSO.",
            "Technical: AuthController, AuthService, JwtTokenService (Nimbus HS256), BCrypt, Flyway users.",
            "Why separate: Security-critical; shortest path to enterprise IdP without redeploying RAG.",
            "Scaling: Token validation should move to gateway JWKS cache; auth DB write-light.",
        ),
        (
            "connector-service (8082)",
            "Business: Integrations (Confluence, Jira, SharePoint); MVP manual upload.",
            "Technical: ManualDocumentConnectorController → IngestionClient WebClient.",
            "Why separate: OAuth credentials and upstream rate limits isolated.",
            "Scaling: Worker pool per source; backpressure when ingestion queue deep.",
        ),
        (
            "ingestion-service (8083)",
            "Business: Normalize, chunk, persist documents.",
            "Technical: DocumentIngestionService @Transactional; TextChunker 800/100; EmbeddingClient.",
            "Why separate: Bursty throughput; should become async via Kafka.",
            "Scaling: Queue consumers; stream large PDFs. Failure: partial index if embedding fails after DB commit.",
        ),
        (
            "embedding-service (8084)",
            "Business: Vector index and similarity search.",
            "Technical: EmbeddingMath hash-based 384-dim vectors; JSON in Postgres; O(N) findAll search.",
            "Why separate: Model upgrades and GPU nodes differ from search API.",
            "Scaling: Replace with pgvector/HNSW; never full table scan in production.",
        ),
        (
            "search-service (8085)",
            "Business: Semantic retrieval API for employees and RAG.",
            "Technical: Thin proxy to embedding search; future ACL pre-filter.",
            "Why separate: Latency SLO distinct from LLM orchestration.",
            "Scaling: Cache hot queries; shard index by tenant.",
        ),
        (
            "rag-service (8086)",
            "Business: Grounded Q&A with citations.",
            "Technical: RagChatService → SearchClient → LlmAnswerService (stub or OpenAI).",
            "Why separate: LLM cost, prompt policy, streaming UX.",
            "Scaling: Dominant cost is LLM tokens; cache cautiously (privacy).",
        ),
        (
            "agent-service (8087) — SCAFFOLD",
            "Planned: Multi-step agents, tool registry, human-in-the-loop.",
            "Why separate: Stateful, long-running, failure-prone vs stateless RAG.",
        ),
        (
            "audit-service (8088) — SCAFFOLD",
            "Planned: Append-only audit_events; DB created, no Flyway yet.",
            "Why separate: Retention and compliance scope differ from auth users.",
        ),
        (
            "notification-service (8089) — SCAFFOLD",
            "Planned: Email/Slack; queue-driven, idempotent delivery.",
        ),
    ]
    for title, *paras in services_detail:
        add_title(doc, title, level=2)
        for p in paras:
            add_para(doc, p)
    doc.add_page_break()

    # --- PHASE 4 ---
    add_title(doc, "Phase 4: Controller / API Analysis", level=1)
    apis = [
        ("POST /api/v1/auth/login", "auth", "Public", "AuthLoginRequest", "AuthTokenResponse", "BCrypt verify; 401 generic message"),
        ("POST /api/v1/auth/register", "auth", "Public", "AuthRegisterRequest", "AuthTokenResponse", "409 if email exists; @Valid 8-128 char password"),
        ("POST /api/v1/connectors/documents", "connector", "JWT", "DocumentUploadRequest", "DocumentIngestionResponse", "X-User-Email from gateway; default sourceType=manual"),
        ("POST /api/v1/ingestion/documents", "ingestion", "Internal*", "DocumentUploadRequest", "DocumentIngestionResponse", "*No JWT on direct port — security gap"),
        ("POST /api/v1/embeddings/index", "embedding", "Internal", "ChunkIndexRequest", "void/200", "Batch index after ingest"),
        ("POST /api/v1/embeddings/search", "embedding", "Internal", "SearchRequest", "SearchResponse", "O(N) full scan — MVP bottleneck"),
        ("POST /api/v1/search/query", "search", "JWT", "SearchRequest", "SearchResponse", "topK 1-20; passthrough to embedding"),
        ("POST /api/v1/rag/chat", "rag", "JWT", "ChatRequest", "ChatResponse", "Chains search + LLM; 5-30s if OpenAI"),
        ("GET /api/v1/health", "all", "Public", "—", "ServiceHealthResponse", "O(1); gateway health is local only"),
    ]
    add_table(
        doc,
        ["Endpoint", "Service", "Auth", "Request", "Response", "Notes"],
        apis,
    )
    add_para(doc, "Request flow (RAG chat):", bold=True)
    add_code(
        doc,
        "Client → LB → Gateway (JWT) → RagChatController → RagChatService.chat\n"
        "  → SearchClient POST /search/query → SearchController → EmbeddingSearchClient\n"
        "  → EmbeddingController → EmbeddingSearchService (DB + cosine)\n"
        "  → LlmAnswerService → OpenAI or stub → ChatResponse",
    )
    doc.add_page_break()

    # --- PHASE 5 ---
    add_title(doc, "Phase 5: Business Logic Analysis", level=1)
    methods = [
        (
            "AuthService.register",
            "Creates user if email unique (case-insensitive), BCrypt hash, issues JWT.",
            "Business: Self-service onboarding.",
            "Alternatives: Invite-only, IdP-only provisioning.",
            "O(1) DB lookups + BCrypt (~100ms). Failure: 409 conflict.",
        ),
        (
            "AuthService.login",
            "Lookup by email, password match, issue JWT.",
            "Business: Sessionless API access.",
            "Alternatives: Refresh tokens, OIDC code flow.",
            "O(1). Failure: 401 (no user enumeration detail in message).",
        ),
        (
            "JwtTokenService.createToken",
            "HS256 JWT with sub, email, name, issuer, exp.",
            "Business: Stateless auth for microservices.",
            "Alternatives: RS256 + JWKS for multi-service verify without shared secret.",
            "O(1). Failure: misconfigured secret breaks entire platform.",
        ),
        (
            "TextChunker.chunk",
            "Sliding window 800 chars, 100 overlap.",
            "Business: Fit embedding model context windows.",
            "Alternatives: Semantic chunking, sentence boundaries.",
            "O(n) in content length. Risk: splits mid-sentence.",
        ),
        (
            "DocumentIngestionService.ingest",
            "Transactional save document + chunks, sync call embedding index.",
            "Business: Make knowledge searchable.",
            "Alternatives: Async Kafka job after persist.",
            "O(chunks). Failure: DB committed but embedding fails → inconsistent index.",
        ),
        (
            "EmbeddingIndexService.index",
            "Per chunk: hash embed, normalize, JSON serialize, save.",
            "Business: Enable similarity search.",
            "Alternatives: OpenAI ada-002, batch API.",
            "O(n) chunks. Not semantically meaningful (dev only).",
        ),
        (
            "EmbeddingSearchService.search",
            "findAll(), embed query, cosine similarity, top-K.",
            "Business: Retrieve relevant chunks.",
            "Alternatives: pgvector IVFFlat/HNSW, dedicated vector DB.",
            "O(N) — first production bottleneck at scale.",
        ),
        (
            "RagChatService.chat",
            "Search then LLM with context.",
            "Business: Grounded answers for employees.",
            "Alternatives: Agentic multi-hop retrieval.",
            "Dominated by search O(N) + LLM latency.",
        ),
        (
            "LlmAnswerService.generate",
            "Build prompt from hits; OpenAI if key else stub quoting excerpts.",
            "Business: Natural language answer with citations.",
            "Alternatives: Azure OpenAI, private LLM, streaming SSE.",
            "Failure: falls back to stub on API error.",
        ),
    ]
    for name, *details in methods:
        add_title(doc, name, level=2)
        labels = ["What", "Why", "Alternatives", "Complexity / failure"]
        for label, text in zip(labels, details):
            add_para(doc, f"{label}: {text}")
    doc.add_page_break()

    # --- PHASE 6 ---
    add_title(doc, "Phase 6: Database Deep Dive", level=1)
    add_title(doc, "users (enterprise_ai_auth)", level=2)
    add_bullets(
        doc,
        [
            "PK: id UUID. Unique email (idx_users_email). password_hash BCrypt.",
            "Why: Separate auth DB for blast radius and backup policy.",
            "Scale: Read replicas for future validation; PgBouncer for connections.",
        ],
    )
    add_title(doc, "documents / document_chunks (enterprise_ai_ingestion)", level=2)
    add_bullets(
        doc,
        [
            "documents: metadata (title, source_type, external_id, owner_email).",
            "document_chunks: FK CASCADE on delete; sequence_no ordering; idx_chunks_document.",
            "Query pattern: load chunks by document_id — index supports.",
            "Sharding: by tenant_id (future column) or source_type.",
        ],
    )
    add_title(doc, "chunk_embeddings (enterprise_ai_embedding)", level=2)
    add_bullets(
        doc,
        [
            "Denormalized: stores document_title and chunk_content for search hits without join.",
            "embedding_json TEXT — production: pgvector column or external index.",
            "Search MVP: SELECT * — no index used for similarity (in-memory).",
            "Replication: read replicas for search; primary for writes.",
        ],
    )
    add_title(doc, "enterprise_ai_audit", level=2)
    add_para(doc, "DB provisioned; no tables. Planned: audit_events partitioned monthly, append-only.")
    doc.add_page_break()

    # --- PHASE 7-9 ---
    add_title(doc, "Phase 7: Kafka / Event Analysis", level=1)
    add_para(doc, "Current state: NO Kafka, RabbitMQ, or Spring Cloud Stream in any pom.xml.")
    add_para(doc, "Why Kafka (planned) vs REST: decouple ingestion latency from API; buffer bursts; replay; audit fan-out.")
    add_table(
        doc,
        ["Topic (planned)", "Producer", "Consumer", "Ordering key"],
        [
            ["ingestion.document.received", "connector", "ingestion", "source_id"],
            ["ingestion.chunk.created", "ingestion", "embedding", "document_id"],
            ["embedding.index.completed", "embedding", "search cache", "—"],
            ["audit.event", "all", "audit", "—"],
        ],
    )
    add_bullets(
        doc,
        [
            "At-least-once + idempotent consumers (document_version UUID).",
            "DLQ: topic.dlq with alert on depth.",
            "Consumer crash: partition reassignment; replay from offset.",
        ],
    )

    add_title(doc, "Phase 8: Cache Analysis", level=1)
    add_para(doc, "Current state: NO Redis or @Cacheable.")
    add_table(
        doc,
        ["Cache (planned)", "Key", "TTL", "Invalidation"],
        [
            ["JWKS", "issuer", "1h", "HTTP cache headers"],
            ["Search results", "tenant:query_hash", "1-5 min", "index.updates event"],
            ["Embeddings", "content_hash:model", "long", "model version bump"],
        ],
    )
    add_bullets(
        doc,
        [
            "Redis down: degrade to origin (higher latency); never cache without tenant in key.",
            "Cache miss: full search/LLM path. Hit ratio target: 30-60% for repeated enterprise FAQs.",
        ],
    )

    add_title(doc, "Phase 9: Security Analysis", level=1)
    add_bullets(
        doc,
        [
            "Auth: POST login/register → HS256 JWT; gateway validates Bearer on all except /auth/** and health.",
            "Authorization: X-User-Email propagated; NO RBAC or document ACL yet.",
            "Risks: shared JWT_SECRET in YAML; internal ports unauthenticated; postgres/postgres dev creds.",
            "PII: email in JWT and owner_email on documents; no encryption at rest in MVP.",
            "Improvements: OIDC (Entra/Okta), RS256+JWKS, mTLS, document ACL at search, audit all queries.",
        ],
    )
    doc.add_page_break()

    # --- PHASE 10-12 ---
    add_title(doc, "Phase 10: Production Readiness Review", level=1)
    add_table(
        doc,
        ["Area", "Current", "Target"],
        [
            ["Logging", "Logback plain text", "JSON + trace_id, tenant_id"],
            ["Metrics", "actuator metrics; no prometheus registry", "micrometer-registry-prometheus"],
            ["Tracing", "None", "OpenTelemetry from gateway"],
            ["Alerts", "None", "5xx rate, DLQ depth, RAG p99, pool exhaustion"],
        ],
    )
    add_title(doc, "Troubleshooting playbooks", level=2)
    troubles = [
        ("Slow API", "Trace gateway→rag→search→embedding; check O(N) embedding search; LLM latency"),
        ("High CPU", "embedding findAll + cosine; scale embedding workers; add vector index"),
        ("High memory", "loading all embeddings in JVM heap per query"),
        ("Kafka lag (future)", "scale consumers; check poison messages in DLQ"),
        ("DB slowness", "Flyway migrations; connection pool; missing indexes on hot paths"),
        ("Redis outage (future)", "cache miss storm; rate limit fail-closed policy"),
    ]
    for issue, fix in troubles:
        add_para(doc, f"{issue}: {fix}")

    add_title(doc, "Phase 11: Scalability Analysis", level=1)
    add_table(
        doc,
        ["Traffic", "First bottleneck", "Second", "Third"],
        [
            ["100 RPS", "Unlikely (LLM cost if all chat)", "—", "—"],
            ["1,000 RPS", "LLM rate limits", "O(N) vector scan", "Gateway CPU"],
            ["10,000 RPS", "Vector index memory", "Audit write volume", "Postgres connections"],
            ["100,000 RPS", "Multi-region index", "LLM spend", "Cross-tenant noisy neighbor"],
            ["1,000,000 RPS", "Clarify: search vs gateway vs telemetry", "Global LB + sharding", "Compliance residency"],
        ],
    )
    add_bullets(
        doc,
        [
            "Horizontal: gateway, search, rag stateless pods.",
            "Vertical: embedding GPU nodes.",
            "DB: PgBouncer, read replicas, pgvector HNSW.",
            "Cache: Redis for hot queries with tenant-scoped keys.",
        ],
    )

    add_title(doc, "Phase 12: Failure Analysis", level=1)
    failures = [
        ("Postgres down", "auth/ingestion/embedding fail startup; no fallback", "Restore replica; RPO per backup policy"),
        ("Kafka down (future)", "Ingestion stops; backlog connectors", "Replay from retained log"),
        ("Redis down (future)", "Higher latency; rate limit policy choice", "Fail closed for abuse protection"),
        ("Service crash", "K8s restart; gateway 503 if all backends down", "HPA, circuit breakers"),
        ("Node crash", "Pod reschedule", "StatefulSet HA for Postgres"),
        ("Network partition", "Split-brain writes risk", "Primary-only writes; stale search disclaimer"),
    ]
    add_table(doc, ["Component", "Customer impact", "Recovery"], failures)
    doc.add_page_break()

    # --- PHASE 13 ---
    add_title(doc, "Phase 13: Interview Preparation", level=1)

    def interview_block(level, q, ideal, mistakes, followups):
        add_title(doc, f"{level}: {q}", level=2)
        add_para(doc, f"Ideal answer: {ideal}")
        add_para(doc, f"Common mistakes: {mistakes}")
        add_para(doc, f"Follow-ups: {followups}")

    interview_block(
        "Senior",
        "Why microservices for this platform?",
        "Different scaling profiles (ingestion vs chat), failure isolation, team autonomy. Trade-off: ops complexity.",
        "Saying microservices are always better; ignoring MVP could be modular monolith first.",
        "When would you merge services? Team <5, low scale.",
    )
    interview_block(
        "Senior",
        "Walk through the MVP RAG path.",
        "Login → upload → ingest chunks → embed → chat calls search → LLM with citations. Mention dev hash embeddings.",
        "Omitting gateway JWT or citing hallucination controls.",
        "How prevent hallucination? Citation-required prompts, confidence thresholds.",
    )
    interview_block(
        "Staff",
        "How secure multi-tenant enterprise knowledge?",
        "Tenant in index keys, ACL on chunks, pre-filter before vector search, audit queries, no cross-tenant cache.",
        "Only app-level JWT without document ACL.",
        "Dedicated vs shared index per tenant?",
    )
    interview_block(
        "Staff",
        "Design ingestion at 10M documents/day.",
        "Kafka pipeline, idempotent workers, horizontal embedding, DLQ, backpressure on connectors.",
        "Synchronous REST chain like MVP.",
        "Exactly-once? At-least-once + idempotent keys.",
    )
    interview_block(
        "Principal",
        "Why separate embedding from search?",
        "GPU batch vs low-latency read, model upgrade cadence, cost accounting.",
        "Cannot articulate when to merge.",
        "At what team size is split not worth it?",
    )
    interview_block(
        "Principal",
        "LLM provider outage strategy?",
        "Multi-provider abstraction, search-only fallback, queue agents, status comms.",
        "No fallback; silent failures.",
        "Regulatory audit requirements?",
    )
    doc.add_page_break()

    # --- PHASE 14 ---
    add_title(doc, "Phase 14: Architecture Critique", level=1)
    critiques = [
        ("Senior Engineer", "Clear boundaries, working MVP slice", "O(N) search, open internal ports", "Extract shared health only if logic grows"),
        ("Staff Engineer", "Monorepo alignment, Flyway on 3 DBs", "No CI/CD, no contract tests", "OpenAPI + Spring Cloud Contract"),
        ("Principal Engineer", "Reference enterprise AI architecture", "10 services ops before PMF", "Validate team size vs boundaries"),
        ("CTO", "Secure enterprise LLM narrative", "Time-to-value vs Python prototypes", "Phase ROI: MVP RAG before agents"),
    ]
    add_table(doc, ["Role", "Strengths", "Weaknesses", "Action"], critiques)

    add_title(doc, "Final Recommendations: Scaling to Millions of Users", level=2)
    add_bullets(
        doc,
        [
            "P0: Replace hash embeddings with OpenAI/Azure + pgvector HNSW; fix O(N) search.",
            "P0: OIDC + RS256 JWT; mTLS service mesh; document ACL at retrieval.",
            "P1: Kafka async ingestion; audit events from RAG and ingest.",
            "P1: Prometheus + distributed tracing; SLO dashboards.",
            "P2: Multi-region read replicas; tenant sharding; LLM gateway with cost caps.",
            "P3: Agent workflows on queue workers; Kubernetes + Helm.",
        ],
    )

    add_title(doc, "Appendix: Complete API & Class Inventory", level=2)
    add_para(
        doc,
        "75 Java source files. Implemented controllers: GatewayHealthController, AuthController, "
        "ManualDocumentConnectorController, DocumentController, EmbeddingController, SearchController, "
        "RagChatController, ServiceHealthController (×9). Services: AuthService, JwtTokenService, "
        "DocumentIngestionService, TextChunker, EmbeddingIndexService, EmbeddingSearchService, EmbeddingMath, "
        "RagChatService, LlmAnswerService. Clients: IngestionClient, EmbeddingClient, EmbeddingSearchClient, SearchClient.",
    )
    add_para(doc, "See repository files: SYSTEM_ARCHITECTURE_ANALYSIS.md, docs/ARCHITECTURE.md, docs/MVP_RUNBOOK.md")

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build_document()
